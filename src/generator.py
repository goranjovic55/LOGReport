from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Preformatted
from docx import Document
from docx.shared import Pt
from typing import List, Dict
from pathlib import Path
import os

from utils.file_utils import filter_lines

class ReportGenerator:
    def __init__(self):
        self.styles = {
            'pdf': {
                'title': ParagraphStyle(
                    'Title',
                    fontName='Helvetica-Bold',
                    fontSize=14,
                    spaceAfter=14,
                    textColor='#5D3E8E'
                ),
                'subtitle': ParagraphStyle(
                    'Subtitle',
                    fontName='Helvetica-Bold',
                    fontSize=12,
                    spaceAfter=6,
                    textColor='#7A5299'
                ),
                'body': ParagraphStyle(
                    'Body',
                    fontName='Courier',
                    fontSize=10,
                    leading=12
                )
            }
        }

    def generate_pdf(self, logs: List[Dict], output_path: str, 
                    lines_mode: str = 'all', 
                    line_limit: int = 0,
                    range_start: int = 0, 
                    range_end: int = 0):
        """Generate PDF with configurable line filtering"""
        try:
            output_path = str(Path(output_path).absolute())
            if not output_path.lower().endswith('.pdf'):
                output_path += '.pdf'
                
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                leftMargin=20*mm,
                rightMargin=20*mm,
                topMargin=20*mm,
                bottomMargin=20*mm
            )
            
            story = []
            for log in logs:
                # Create a copy of the log to modify
                processed_log = dict(log)
                
                # Apply line filtering based on parameters
                processed_log['content'] = filter_lines(
                    processed_log['content'],
                    mode=lines_mode,
                    limit=line_limit,
                    start=range_start,
                    end=range_end
                )
                
                # Add section
                story.append(Paragraph(
                    f"File: {processed_log['filename']}",
                    self.styles['pdf']['title']
                ))
                
                # Add filtered content as preformatted text to preserve whitespace
                content_text = "\n".join(processed_log['content'])
                story.append(Preformatted(content_text, self.styles['pdf']['body']))
                
                story.append(Spacer(1, 12*mm))
            
            doc.build(story)

        except Exception as e:
            raise RuntimeError(f"PDF generation failed: {str(e)}")

    def generate_docx(self, logs: List[Dict], output_path: str, lines_mode: str = 'all'):
        """Generate DOCX with line filtering"""
        doc = Document()
        
        # Configure styles
        styles = doc.styles
        title_style = styles.add_style('LogTitle', 1)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(14)
        title_style.font.color.rgb = 0x5D3E8E
        
        body_style = styles['Normal']
        body_style.font.name = 'Courier New'
        body_style.font.size = Pt(10)
        
        for log in logs:
            # Add title
            doc.add_paragraph(log['filename'], style='LogTitle')
            
            # Process lines
            content = filter_lines(
                log['content'],
                mode=lines_mode,
                limit=log.get('line_limit', 0),
                start=log.get('range_start', 0),
                end=log.get('range_end', 0)
            )
            
            # Combine all lines into a single paragraph with preserved newlines
            full_text = '\n'.join(content)
            doc.add_paragraph(full_text, style='Normal')
            
            doc.add_page_break()
            
        doc.save(output_path)